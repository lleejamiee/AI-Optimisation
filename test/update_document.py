import os
from dotenv import load_dotenv
from llama_index.core import Settings
from llama_index.llms.groq import Groq
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_parse import LlamaParse
from haystack.components.preprocessors import DocumentSplitter
from haystack import Document
from llama_index.core import VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.retrievers import RecursiveRetriever
from llama_index.core.schema import TextNode
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core import QueryBundle
from llama_index.core.postprocessor import LLMRerank

from autogen import ConversableAgent

from docx import Document

doc = Document("../uploaded_files/outdated_FPH ICT Department - New Graduate Onboarding and Training Guide.docx")

old_text = ("Slack: While Microsoft Teams is the primary tool \n"
            "for meetings and chats, Slack is occasionally used for quick, \n"
            "informal communication between smaller project teams.")
new_text = "Microsoft Teams: Microsoft Team is the primary tool for meetings, chats, as well as quick and informal communication between smaller project teams."

print(old_text)

for paragraph in doc.paragraphs:
    full_text = "".join([run.text for run in paragraph.runs])

    if old_text in paragraph.text:
        new_full_text = full_text.replace(old_text, new_text)

        for run in paragraph.runs:
            run.clear()

        paragraph.add_run(new_full_text)

doc.save("../uploaded_files/updated.docx")

# load_dotenv()
#
# Settings.llm = Groq(model="llama-3.1-70b-versatile", api_key=os.getenv("GROQ_API_KEY"), temperature=0)
# Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
#
# outdated_docs = LlamaParse(num_workers=8, split_by_page=0, result_type="text").load_data(
#     "../uploaded_files/outdated_FPH ICT Department - New Graduate Onboarding and Training Guide.docx")
# reference_docs = LlamaParse(split_by_page=0, result_type="text").load_data("../uploaded_files/instructions.docx")
#
# # Splitter
# outdated_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=90)
# reference_splitter = DocumentSplitter(split_by="passage", split_length=1, split_overlap=0)
#
# # Get Chunks
# outdated_chunks = outdated_splitter.get_nodes_from_documents(outdated_docs)
# reference_chunks = reference_splitter.run(documents=[Document(content=reference_docs[0].text)])
#
#
# def split_into_paragraphs(text):
#     paragraphs = text.split('\n\n')
#     return [para.strip() for para in paragraphs if para.strip()]
#
#
# def create_subnodes(outdated_chunks):
#     all_subnodes = []
#     subnode_to_basenode_map = {}
#
#     for base_node in outdated_chunks:
#         paragraphs = split_into_paragraphs(base_node.text)
#
#         for paragraph in paragraphs:
#             sub_node = TextNode(text=paragraph, metadata=base_node.metadata.copy())
#             sub_node.metadata["base_node_id"] = base_node.node_id
#             all_subnodes.append(sub_node)
#             subnode_to_basenode_map[sub_node.node_id] = base_node
#
#     return all_subnodes, subnode_to_basenode_map
#
#
# def retrieve_subnodes(subnodes, nodes_map, query, vector_store_index, top_k=2):
#     vector_retriever = vector_store_index.as_retriever(similarity_top_k=top_k)
#
#     all_nodes_dict = {n.node_id: n for n in subnodes}
#
#     retriever = RecursiveRetriever(
#         "vector",
#         retriever_dict={"vector": vector_retriever},
#         node_dict=all_nodes_dict,
#         verbose=True,
#     )
#
#     retrieved_nodes = retriever.retrieve(query)
#
#     best_node = retrieved_nodes[0]
#     highest_scoring_base_node = nodes_map[best_node.node.node_id]
#
#     # return only best node instead of list[retrieved_nodes]
#     return retrieved_nodes, highest_scoring_base_node
#
#
# all_subnodes, subnode_to_basenode_map = create_subnodes(outdated_chunks)
# vector_index = VectorStoreIndex(all_subnodes)
#
# queries = []
#
# for chunk in reference_chunks['documents']:
#     queries.append(chunk.content)
#
#     for query in queries:
#         retrieved_subnodes, top_base_node = retrieve_subnodes(all_subnodes, subnode_to_basenode_map, query,
#                                                               vector_index)
#
#         # print(f"Score: {retrieved_subnodes[0].score}\n\n{retrieved_subnodes[0].text}")
#         for subnode in retrieved_subnodes:
#             print(subnode.text)
#             print("-------" * 40)
#
#         print(top_base_node.text)
#
# nodes = []
# for chunk in outdated_chunks:
#     # nodes.append(TextNode(text=chunk.content, id_=chunk.meta['id']))
#     nodes.append(TextNode(text=chunk.get_content()))
#
# outdated_index = VectorStoreIndex(nodes)
#
#
# def get_retrieved_nodes(
#         query_str, vector_top_k=10, reranker_top_n=3, with_reranker=False
# ):
#     query_bundle = QueryBundle(query_str)
#     # configure retriever
#     retriever = VectorIndexRetriever(
#         index=outdated_index,
#         similarity_top_k=vector_top_k,
#     )
#     retrieved_nodes = retriever.retrieve(query_bundle)
#
#     if with_reranker:
#         # configure reranker
#         reranker = LLMRerank(
#             choice_batch_size=5,
#             top_n=reranker_top_n,
#         )
#         retrieved_nodes = reranker.postprocess_nodes(
#             retrieved_nodes, query_bundle
#         )
#
#     return retrieved_nodes
#
#
# # Create agents, takes a list of the vector embeddings for the refernece chunks and the vector embeddings for the outdated chunks
# def create_pairs(reference):
#     context_pairs = []
#     # result_nodes = []
#     # retriever = outdated_index.as_retriever()
#     i = 0
#     for text_chunk in reference:
#         print(i)
#         result_nodes = get_retrieved_nodes(
#             text_chunk.content
#         )
#
#         # Save result as a Tuple
#         result = (text_chunk, result_nodes[0])
#
#         # Append result to a List
#         context_pairs.append(result)
#         i += 1
#
#     return context_pairs
#
#
# context_pairs = create_pairs(reference_chunks['documents'])
#
# print("Context pairs")
# print(context_pairs[0][0].content)
# print(context_pairs[0][1].text)
#
# config_list = [
#     {
#         "model": "llama-3.1-70b-versatile",
#         "api_key": os.environ.get("GROQ_API_KEY"),
#         "api_type": "groq",
#         "frequency_penalty": 0.5,
#         "max_tokens": 8000,
#         "presence_penalty": 0.2,
#         "seed": 42,
#         "temperature": 0,
#         "top_p": 0.2
#     }
# ]
#
# agent = ConversableAgent(
#     "chatbot",
#     llm_config={"config_list": config_list},
#     code_execution_config=False,  # Turn off code execution, by default it is off.
#     function_map=None,  # No registered functions, by default it is None.
#     human_input_mode="NEVER",  # Never ask for human input.
# )
#
# # This 1 for sys_prompt
# sys_prompt = """
# You are tasked with identifying and retaining only the sections of the document that need to be replaced based on the provided reference material. Your goal is to extract the content that is directly relevant for replacement without adding or modifying the existing text.
#
# Instructions:
# 1. Review both the provided document chunk and reference material to determine which sections are directly relevant for replacement.
# 2. Extract the content that will be replaced, including any references or context that should be retained as part of the replacement.
# 3. Ensure that the extracted section remains as close as possible in length and context to the original content that will be replaced.
# 4. Do not include any additional content or modifications outside of what is necessary for the update.
# 5. Do not include any explanations, introductions, or additional text.
# """
#
# # This 1 for context
# context = f"""
# Please extract and retain only the section of the document that is directly relevant to the content being updated. Here is the specific section to keep:
#
# Here is a chunk from a report that needs refinement:
# {context_pairs[0][1].text}
#
# Here is the reference material to use for determining relevance:
# {context_pairs[0][0].content}
#
# Ensure that only this section is extracted and preserved for the update.
# """
#
# # Define the system message and the user's prompt
# system_message = {"content": sys_prompt, "role": "system"}
#
# # Define the user's message
# user_message = {"content": context, "role": "user"}
#
# # Generate a reply using the system and user messages
# reply = agent.generate_reply(messages=[system_message, user_message])
#
# # Output the reply
# print("Reply from agent")
# print("------------------------------------------------")
# print(reply['content'])
#
# """
# Instead of hard coding line 218 and 221, make them dynamic. As we loop through the context pair, we make another pair:
# relevant part, reference
# We loop through this pair, we feed them into llm and generate a pair for update:
# relevant part, update content
# We use the code on the top to replace text.
# """