import os
from pathlib import Path
import re
from docx import Document
from llama_index.core import Settings, VectorStoreIndex, Document as LlamaDocument
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from llama_parse import LlamaParse
from llama_index.core.tools import RetrieverTool
from llama_index.core.retrievers import VectorIndexRetriever, RouterRetriever
from llama_index.retrievers.bm25 import BM25Retriever
from llama_index.core.node_parser import SentenceSplitter
from utilities.prompts import update_docs_system
import pdf2docx
from dotenv import load_dotenv

load_dotenv()


class DocumentRetriever:
    def __init__(self):
        Settings.llm = Groq(model="llama-3.2-90b-text-preview", api_key=os.getenv("GROQ_API_KEY"), temperature=0)
        Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")

        os.getenv("LLAMA_CLOUD_API_KEY")
        os.getenv("GROQ_API_KEY")

    # Process reference material and split into paragraphs
    def _process_reference(self, file_path):
        reference_file = self._save_uploaded_file(file_path)
        reference = LlamaParse(num_workers=8, split_by_page=0, result_type="text").load_data(reference_file)
        reference_split = self._split_into_paragraphs(reference[0].text)

        return reference_split

    # Save uploaded file to a local directory
    def _save_uploaded_file(self, file):
        save_folder = Path('uploaded_files/')
        save_folder.mkdir(parents=True, exist_ok=True)
        save_path = save_folder / file.name

        with open(save_path, mode='wb') as w:
            w.write(file.getvalue())

        return save_path

    # Split text into paragraphs for reference material
    def _split_into_paragraphs(self, text):
        paragraphs = text.split('\n\n')
        return [para.strip() for para in paragraphs if para.strip()]

    # Split document text into sections based on headings
    def _process_document(self, file_path):
        saved_file = self._save_uploaded_file(file_path)
        extension = saved_file.suffix.lower()

        if extension == ".pdf":
            converted = self._convert_to_docx(saved_file)
            document = Document(converted)
        else:
            document = Document(saved_file)
        doc_text = "\n".join([para.text for para in document.paragraphs])

        return self._split_into_sections(doc_text)

    # Convert PDF to DOCX format
    def _convert_to_docx(self, pdf_file):
        docx_file = pdf_file.with_suffix(".docx")
        pdf2docx.parse(str(pdf_file), str(docx_file))

        return docx_file

    # Split document text into sections based on headings
    def _split_into_sections(self, doc_text):
        sections = []

        pattern = re.compile(r'''
            ^
            (?P<Section>\d+(?:\.\d+)*)\s+
            (?P<Title>.+?)
            \n
            (?P<Content>(?:(?!^\d+(?:\.\d+)*\s+).+\n?)+)
        ''', re.MULTILINE | re.VERBOSE)

        for match in pattern.finditer(doc_text):
            sections.append({
                "title": f"{match.group('Section')} {match.group('Title')}",
                "content": match.group('Content').strip()
            })

        return sections

    # Create objects from sections
    def _create_documents(self, sections):
        documents = []
        for section in sections:
            text = f"{section['title']}\n\n{section['content']}"
            doc = LlamaDocument(text=text)
            documents.append(doc)

        return documents

    # Create vector index with documents
    def _create_index(self, documents, vector_index=None):
        if vector_index is None:
            vector_index = VectorStoreIndex([])
        for doc in documents:
            vector_index.insert(doc)

        return vector_index

    # Create a router retriever with vector and BM25 retrievers
    def _create_retriever(self, documents, vector_index):
        parser = SentenceSplitter()
        nodes = parser.get_nodes_from_documents(documents)
        vector_retriever = VectorIndexRetriever(vector_index, similarity_top_k=6)
        bm25_retriever = BM25Retriever.from_defaults(nodes=nodes, similarity_top_k=6)

        retriever_tools = [
            RetrieverTool.from_defaults(
                retriever=vector_retriever,
                description="Useful in most cases",
            ),
            RetrieverTool.from_defaults(
                retriever=bm25_retriever,
                description="Useful if searching about specific information",
            ),
        ]

        router_retriever = RouterRetriever.from_defaults(
            retriever_tools=retriever_tools,
            select_multi=True,
        )

        return router_retriever

    # RAG pipeline
    def _rag_components(self, file_path):
        sections = self._process_document(file_path)
        documents = self._create_documents(sections)
        vector_index = self._create_index(documents)
        retriever = self._create_retriever(documents, vector_index)

        return retriever

    # Perform retrieval using reference material and additional entries
    def rag_retrieval(self, outdated, reference_material=None, additional_search_text=None):
        context_pairs = []
        search_text = []

        retriever = self._rag_components(outdated)

        if reference_material is not None:
            search_text.extend(self._process_reference(reference_material))
        if additional_search_text is not None:
            search_text.extend(additional_search_text)

        for reference in search_text:
            retrieved = retriever.retrieve(reference)
            context_pairs.append({'reference': reference, 'retrieved': retrieved})

        return context_pairs

    # Update the sections of the document using an LLM with provided context
    def update_sections(self, sections_to_update):
        updated_sections = []
        for section in sections_to_update:
            reference = section['reference']
            retrieved = section['retrieved']

            # Construct the prompt by formatting the update_docs_system string
            prompt = update_docs_system.format(reference=reference, retrieved=retrieved)

            # Use the LLM to generate the updated content
            response = Settings.llm.complete(prompt)
            updated_content = response.text.strip()

            updated_sections.append({
                'reference': reference,
                'original_content': retrieved.text,
                'updated_content': updated_content
            })

        return updated_sections